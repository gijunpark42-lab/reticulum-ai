# write_research_log.py — regenerates RESEARCH_LOG.docx from quant/results.json.
#
# Document structure (one Word file for ALL studies, per the owner's decision):
#   1. Universe (the 71 calls, listed in full)
#   2. Common methodology (incl. benchmarks 1-3 = SOXX/SPY/QQQ and
#      benchmarks 4-6 = the universe-average abnormal return per index)
#   3. Study 1: full-capacity binary signal  (future signals append as Study 2, 3, ...)
#   N. Update history (every docx write, with date+time — quant/log_history.csv)
#
# HOW TO RUN (from the project root; close the docx in Word first)
#   python -X utf8 quant/write_research_log.py "note describing what changed"

import json
import csv
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt

RESULTS_FILE = "quant/results.json"
SIGNAL_FILE = "quant/signal_full_capacity.csv"
CALLS_FILE = "quant/calls.csv"
HISTORY_FILE = "quant/log_history.csv"
OUT_FILE = "RESEARCH_LOG.docx"

BENCHMARKS = ["SOXX", "SPY", "QQQ"]
WINDOWS = [5, 10, 20]


def pct(x):
    return f"{x*100:+.1f}%" if x is not None else "—"


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = str(c)
        cell.paragraphs[0].runs[0].bold = True
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = str(v)


def append_history(note):
    new = not os.path.exists(HISTORY_FILE)
    with open(HISTORY_FILE, "a", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        if new:
            writer.writerow(["written_at", "note"])
        writer.writerow([datetime.now().strftime("%Y-%m-%d %H:%M"), note])


def load_history():
    if not os.path.exists(HISTORY_FILE):
        return []
    with open(HISTORY_FILE, encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def group_results_table(doc, results, suffix, groups=("full_capacity", "rest"),
                        labels=("full capacity (flag=1)", "rest (flag=0)")):
    """One results table for a metric suffix (e.g. 'SOXX' or 'SOXX_minus_univ')."""
    rows = []
    for g, glabel in zip(groups, labels):
        for w in WINDOWS:
            s = results["groups"][g].get(f"drift{w}_vs_{suffix}")
            if s:
                rows.append([glabel, f"+{w}d", s["n"], pct(s["mean"]), pct(s["median"]),
                             f"{s['hit_rate']*100:.0f}%",
                             s["p_value"] if s["p_value"] is not None else "—"])
    add_table(doc, ["Group", "Window", "N", "Mean", "Median", "Hit rate", "p-value"], rows)
    doc.add_paragraph()


def main():
    note = sys.argv[1] if len(sys.argv) > 1 else "log regenerated"
    append_history(note)
    history = load_history()

    with open(RESULTS_FILE, encoding="utf-8") as f:
        results = json.load(f)
    with open(SIGNAL_FILE, encoding="utf-8-sig") as f:
        signal_rows = list(csv.DictReader(f))
    with open(CALLS_FILE, encoding="utf-8-sig") as f:
        calls = list(csv.DictReader(f))
    calls.sort(key=lambda c: c["company"].lower())
    dates = sorted(c["date"] for c in calls)
    flagged = {r["company"] for r in signal_rows if r["flag"] == "1"}

    doc = Document()
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading("Reticulum AI — Quant Research Log", level=0)
    doc.add_paragraph(f"Last updated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  ·  "
                      f"AI/semiconductor supply-chain graph (177 nodes / 997 edges)")

    # ── 1. Universe ──────────────────────────────────────────────────────────
    doc.add_heading("1. Universe — all earnings calls in the study", level=1)
    doc.add_paragraph(
        f"{len(calls)} US-listed earnings calls (one per company), {dates[0]} – {dates[-1]}, "
        "drawn from the hand-curated supply-chain graph. Chain signals only — raw "
        "transcripts, event keynotes (GTC) and news/research notes are not used. "
        "A call counts for a company only when the source label is that company's own call."
    )
    add_table(doc, ["Company", "Ticker", "Call date"],
              [[c["company"], c["ticker"], c["date"]] for c in calls])
    doc.add_paragraph()

    # ── 2. Common methodology ────────────────────────────────────────────────
    doc.add_heading("2. Common methodology (applies to every study)", level=1)
    for line in [
        "t0 = the earnings-call date.",
        "Entry = the CLOSE of t+1 (first trading day after the call). The announcement "
        "jump close(t0)→close(t+1) is not buyable and is reported separately as context.",
        "Windows = entry → +5, +10, +20 trading days. A recent call missing +20d data "
        "still contributes its +5d/+10d windows (membership: quant/window_companies.csv).",
        "Benchmarks 1–3: abnormal return = stock return minus index return over the "
        "same window, for SOXX (semiconductor sector), SPY (market), QQQ (large-cap tech).",
        "Benchmarks 4–6: the UNIVERSE AVERAGE — for each window/index, the mean abnormal "
        "return across all calls in the universe. A group's excess = its abnormal return "
        "minus this average. This asks directly: did the signal group beat the average "
        "call in this universe?",
        "Statistics per group and window: N, mean, median, hit rate, two-sided p-value "
        "of a one-sample t-test of the mean against zero (p < 0.05 conventional).",
        "Classification governance: signals are BINARY and owner-approved per company; "
        "candidates come from chain signals reviewed in batches of 10 chain files, and "
        "ambiguous cases are individually ruled allow/deny by the project owner.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── 3. Study 1: full-capacity binary signal ──────────────────────────────
    doc.add_heading("3. Study 1 — Full-capacity binary signal", level=1)
    doc.add_paragraph(
        "Question: when a company states on its own earnings call that its capacity is "
        "full (sold out / fully allocated / demand exceeds its supply) — regardless of "
        "any input-material narrative — does its stock subsequently earn abnormal "
        "returns versus benchmarks 1–6?"
    )
    doc.add_heading(f"Flagged companies (flag = 1): {len(flagged)} of {len(calls)}", level=2)
    add_table(doc, ["Company", "Evidence (chain signal)", "Source", "Ruling"],
              [[r["company"], r["evidence"], r["source"], r["owner_ruling"]]
               for r in signal_rows if r["flag"] == "1"])
    doc.add_paragraph()

    excl = results.get("excluded", {})
    if excl:
        doc.add_paragraph(
            "Window exclusions (back-fill automatically as time passes): " +
            "; ".join(f"{k}: {v}" for k, v in excl.items()) + "."
        )
    doc.add_heading("Results — benchmarks 1–3 (index-relative)", level=2)
    for bench in BENCHMARKS:
        doc.add_heading(f"Abnormal returns vs {bench}", level=3)
        group_results_table(doc, results, bench)
    doc.add_heading("Results — benchmarks 4–6 (excess over the universe average)", level=2)
    for bench in BENCHMARKS:
        doc.add_heading(f"Excess over universe-average abnormal-vs-{bench}", level=3)
        group_results_table(doc, results, f"{bench}_minus_univ")

    doc.add_heading("Reading of the results", level=2)
    for line in [
        "The flagged group underperforms on every benchmark and window. The single "
        "conventionally significant cell is +20d vs SOXX: mean −9.4%, median −13.2%, "
        "hit rate 13%, p = 0.026.",
        "Against benchmarks 4–6 (the universe average), the flagged group is below "
        "average in every cell (−2% to −6%) but no cell clears p < 0.05 "
        "(best p = 0.10, +5d vs SPY-average) — directionally consistent underperformance, "
        "not yet statistically significant at this sample size.",
        "Versus SPY and QQQ the REST group shows strong positive drift (+14.2% / +10.2% "
        "at +20d, p ≤ 0.005) while the flagged group is roughly flat — saying 'we are "
        "sold out' was associated with LAGGING the rest of the supply chain.",
        "Candidate interpretations (untested): (1) a public sell-out statement marks "
        "peak expectations and is already priced; (2) full capacity is a growth ceiling "
        "while unconstrained names can still surprise; (3) single-period artifact of "
        "the Feb–Jun 2026 sample.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Limitations", level=2)
    for line in [
        "Small sample: 17 flagged vs 54 rest, one 4-month window, one market regime; "
        "April–May call clustering overlaps the measurement windows.",
        "Retroactive collection: t0 uses the public call date (no look-ahead in "
        "measurement) but no live-trading claim is made.",
        "Binary flag ignores intensity (TSMC 'very tight' = Fabrinet 'constrained' = 1).",
        "Universe selection tilts toward AI-demand names by construction of the graph.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    # ── Update history ───────────────────────────────────────────────────────
    doc.add_heading("Update history", level=1)
    add_table(doc, ["Written at", "Note"],
              [[h["written_at"], h["note"]] for h in history])

    doc.save(OUT_FILE)
    print(f"Wrote {OUT_FILE} ({history[-1]['written_at']})")


if __name__ == "__main__":
    main()
